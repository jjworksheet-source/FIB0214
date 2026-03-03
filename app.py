import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
import pandas as pd
import datetime
import io
import os
import re
import base64
import random
import time

from pdf2image import convert_from_bytes
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail, Attachment, FileContent, FileName, FileType, Disposition, Email
from python_http_client.exceptions import HTTPError

# ============================================================
# --- Streamlit Setup ---
# ============================================================

st.set_page_config(page_title="Worksheet Generator", page_icon="📝", layout="wide")
st.title("📝 校本填充工作紙生成器")

# Session state
st.session_state.setdefault("shuffled_cache", {})
st.session_state.setdefault("final_pool", {})
st.session_state.setdefault("ai_choices", {})
st.session_state.setdefault("confirmed_batches", set())
st.session_state.setdefault("last_selected_level", None)
st.session_state.setdefault("selected_student_name_b", None)

# 防止 final_pool 被污染
if not isinstance(st.session_state.final_pool, dict):
    st.session_state.final_pool = {}

# ============================================================
# --- ReportLab Font Setup ---
# ============================================================

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_paths = [
        "Kai.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"
    ]

    CHINESE_FONT = None
    for path in font_paths:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("ChineseFont", path))
                CHINESE_FONT = "ChineseFont"
                break
            except Exception:
                continue

    if not CHINESE_FONT:
        st.error("❌ Chinese font not found. Please ensure Kai.ttf is in your GitHub repository.")

except ImportError:
    st.error("❌ reportlab not found. Please add 'reportlab' to your requirements.txt")
    st.stop()

# ============================================================
# --- Google Sheet Connection ---
# ============================================================

try:
    key_dict = st.secrets["gcp_service_account"]
    creds = Credentials.from_service_account_info(
        key_dict,
        scopes=[
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive.file"
        ]
    )
    client = gspread.authorize(creds)
    SHEET_ID = st.secrets["app_config"]["spreadsheet_id"]

except Exception as e:
    st.error(f"❌ Google Sheet Connection Error: {e}")
    st.stop()

# ============================================================
# --- Google Sheet Loader ---
# ============================================================

def load_sheet(sheet_name: str) -> pd.DataFrame:
    try:
        sh = client.open_by_key(SHEET_ID)
        ws = sh.worksheet(sheet_name)
        df = pd.DataFrame(ws.get_all_records())
        df.columns = [c.strip() for c in df.columns]
        for col in df.columns:
            df[col] = df[col].astype(str).str.strip()
        return df
    except Exception as e:
        st.error(f"❌ 無法讀取工作表「{sheet_name}」: {e}")
        return pd.DataFrame()


@st.cache_data(ttl=60)
def load_students():
    return load_sheet("學生資料")


@st.cache_data(ttl=60)
def load_standby():
    """載入 standby 工作表（題庫）"""
    try:
        return load_sheet("standby")
    except Exception:
        return pd.DataFrame()


def update_status_to_used(row_indices):
    """更新 standby 工作表中句子的狀態為已使用"""
    try:
        sh = client.open_by_key(SHEET_ID)
        ws = sh.worksheet("standby")
        for idx in row_indices:
            gs_row = idx + 2  # pandas 0-based → Google Sheets 1-based (header = row 1)
            ws.update_cell(gs_row, 8, "已使用")  # Status 是第 8 欄
        return True, f"成功更新 {len(row_indices)} 筆記錄"
    except Exception as e:
        return False, str(e)

# ============================================================
# --- standby Parser ---
# ============================================================

def parse_standby_table(df: pd.DataFrame):
    """
    解析 standby 表格
    欄位：ID, School, level, Word, Type, Content, Answer, Status, Entry_Date
    跳過 Status 為「已使用」的句子
    """
    groups = {}

    for idx, row in df.iterrows():
        school  = str(row.get("School", "")).strip()
        level   = str(row.get("level", "")).strip()   # 小寫 level
        word    = str(row.get("Word", "")).strip()
        content = str(row.get("Content", "")).strip()
        status  = str(row.get("Status", "")).strip()

        if not (school and level and word and content):
            continue
        if status == "已使用":
            continue

        batch_key = f"{school}||{level}"
        groups.setdefault(batch_key, {})
        groups[batch_key].setdefault(word, {
            "content": content,
            "is_ready": True,
            "row_index": idx
        })

    return groups

# ============================================================
# --- Batch Readiness Checker ---
# ============================================================

def compute_batch_readiness(batch_key: str, word_dict: dict):
    """所有句子都已就緒（standby 已預先審核）"""
    ready_words = []
    for word, data in word_dict.items():
        if data.get("is_ready") and data.get("content"):
            ready_words.append((word, data["content"]))
    return ready_words, [], True

# ============================================================
# --- Final Pool Builder ---
# ============================================================

def build_final_pool_for_batch(batch_key: str, word_dict: dict):
    """直接使用 standby 中所有可用句子"""
    school, level = batch_key.split("||")
    questions = []
    for word, data in word_dict.items():
        content = data.get("content", "")
        if content:
            questions.append({
                "Word": word,
                "Content": content,
                "School": school,
                "Level": level,
            })
    return questions

# ============================================================
# --- PDF Text Rendering Helpers ---
# ============================================================

def draw_text_with_underline_wrapped(c, x, y, text, font_name, font_size, max_width,
                                     underline_offset=2, line_height=18):
    parts = re.split(r'(<u>.*?</u>)', text)
    tokens = []

    for p in parts:
        if not p:
            continue
        if p.startswith("<u>") and p.endswith("</u>"):
            tokens.append(p)
        else:
            tokens.extend(list(p))

    def measure(tok):
        inner = tok[3:-4] if tok.startswith("<u>") else tok
        return pdfmetrics.stringWidth(inner, font_name, font_size)

    def draw_line(line_tokens, draw_x, draw_y):
        cx = draw_x
        for tp in line_tokens:
            c.setFont(font_name, font_size)
            if tp.startswith("<u>") and tp.endswith("</u>"):
                inner = tp[3:-4]
                c.drawString(cx, draw_y, inner)
                w = pdfmetrics.stringWidth(inner, font_name, font_size)
                c.line(cx, draw_y - underline_offset, cx + w, draw_y - underline_offset)
                cx += w
            else:
                c.drawString(cx, draw_y, tp)
                cx += pdfmetrics.stringWidth(tp, font_name, font_size)

    cur_y = y
    line_buf, line_width = [], 0

    for tok in tokens:
        tok_w = measure(tok)
        if line_width + tok_w > max_width and line_buf:
            draw_line(line_buf, x, cur_y)
            cur_y -= line_height
            line_buf, line_width = [tok], tok_w
        else:
            line_buf.append(tok)
            line_width += tok_w

    if line_buf:
        draw_line(line_buf, x, cur_y)
        cur_y -= line_height

    cur_y -= 12
    return cur_y

# ============================================================
# --- Student Worksheet PDF Generator (WITH HEADER ON EVERY PAGE) ---
# ============================================================

def create_pdf(school_name, level, questions, student_name=None):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Frame, PageTemplate
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter

    bio = io.BytesIO()

    # --- CUSTOM PAGE TEMPLATE WITH HEADER ---
    def header_footer(canvas, doc):
        canvas.saveState()
        font_name = CHINESE_FONT if CHINESE_FONT else 'Helvetica'
        canvas.setFont(font_name, 23)
        canvas.drawCentredString(letter[0] / 2, letter[1] - 1*inch, "童學童樂教育中心")
        canvas.restoreState()

    frame = Frame(0.75*inch, 0.75*inch, letter[0]-1.5*inch, letter[1]-2*inch, id='normal')
    template = PageTemplate(id='header_template', frames=frame, onPage=header_footer)
    doc = SimpleDocTemplate(bio, pagesize=letter)
    doc.addPageTemplates(template)

    story = []

    styles = getSampleStyleSheet()
    font_name = CHINESE_FONT if CHINESE_FONT else 'Helvetica'

    # --- STYLES ---
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=22,
        alignment=TA_CENTER,
        spaceAfter=12
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=18,
        leading=26,
        leftIndent=0,
        firstLineIndent=0
    )

    vocab_title_style = ParagraphStyle(
        'VocabTitle',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=20,
        alignment=TA_CENTER,
        spaceAfter=20
    )

    # --- TITLE & DATE (AFTER HEADER) ---
    title_text = f"<b>{school_name} ({level}) - {student_name} - 校本填充工作紙</b>" if student_name \
                 else f"<b>{school_name} ({level}) - 校本填充工作紙</b>"
    story.append(Paragraph(title_text, title_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"日期: {datetime.date.today() + datetime.timedelta(days=1)}", normal_style))
    story.append(Spacer(1, 0.3*inch))

    # --- QUESTIONS ---
    for i, row in enumerate(questions):
        content = row['Content']
        content = re.sub(r'【】(.*?)【】', r'<u>\1</u>', content)
        content = re.sub(r'【(.+?)】', r'<u>________</u>', content)

        num_para = Paragraph(f"<b>{i+1}.</b>", normal_style)
        content_para = Paragraph(content, normal_style)

        t = Table([[num_para, content_para]], colWidths=[0.5*inch, 6.7*inch])
        t.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.15*inch))

    # --- VOCABULARY TABLE (Second Page) ---
    # 1. Extract words in the same order as the questions
    words = [row.get('Word', '').strip() for row in questions]
    
    # 2. Remove duplicates while preserving order (Python 3.7+ dict behavior)
    unique_words = list(dict.fromkeys([w for w in words if w]))

    if unique_words:
        story.append(PageBreak())
        story.append(Paragraph("<b>詞語表</b>", vocab_title_style))
        story.append(Spacer(1, 0.2*inch))

        num_cols = 4
        table_data = []
        for i in range(0, len(unique_words), num_cols):
            row = unique_words[i:i+num_cols]
            while len(row) < num_cols:
                row.append('')
            table_data.append(row)

        col_width = 1.8*inch
        vocab_table = Table(table_data, colWidths=[col_width]*num_cols)
        vocab_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 22),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('TOPPADDING', (0, 0), (-1, -1), 16),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 16),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(vocab_table)

    doc.build(story)
    bio.seek(0)
    return bio

# ============================================================
# --- Teacher Answer PDF Generator ---
# ============================================================

def create_answer_pdf(school_name, level, questions):
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.colors import red as RED

    bio = io.BytesIO()
    c = rl_canvas.Canvas(bio, pagesize=letter)
    page_width, page_height = letter
    font_name = CHINESE_FONT or "Helvetica"

    cur_y = page_height - 80
    left_m = 60

    c.setFont(font_name, 22)
    c.drawString(left_m, cur_y, "詞語清單（題目順序）")
    cur_y -= 40

    c.setFont(font_name, 18)

    for idx, row in enumerate(questions, start=1):
        word = row["Word"]

        if cur_y < 60:
            c.showPage()
            cur_y = page_height - 80
            c.setFont(font_name, 22)
            c.drawString(left_m, cur_y, "詞語清單（續）")
            cur_y -= 40
            c.setFont(font_name, 18)

        c.drawString(left_m, cur_y, f"{idx}. ")
        c.setFillColor(RED)
        c.drawString(left_m + 40, cur_y, word)
        c.setFillColorRGB(0, 0, 0)
        cur_y -= 26

    c.save()
    bio.seek(0)
    return bio

# ============================================================
# --- DOCX Worksheet Generator ---
# ============================================================

def create_docx(school_name, level, questions, student_name=None):
    doc = Document()

    title_text = f"{school_name} ({level}) - {student_name} - 校本填充工作紙" if student_name \
                 else f"{school_name} ({level}) - 校本填充工作紙"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f"日期: {datetime.date.today() + datetime.timedelta(days=1)}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")

    for i, row in enumerate(questions):
        content = re.sub(r'【|】', '', row["Content"])
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(content)
        run.font.size = Pt(18)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ============================================================
# --- SendGrid Email Sender ---
# ============================================================

def send_email_with_pdf(to_email, student_name, school_name, grade, pdf_bytes, cc_email=None):
    try:
        sg_config = st.secrets["sendgrid"]
        recipient = str(to_email).strip()

        if not re.match(r'^[\w\.-]+@[\w\.-]+\.\w+$', recipient):
            return False, f"無效的家長電郵格式: '{recipient}'"

        from_email_obj = Email(sg_config["from_email"], sg_config.get("from_name", ""))
        safe_name = re.sub(r'[^\w\-]', '_', str(student_name).strip())

        message = Mail(
            from_email=from_email_obj,
            to_emails=recipient,
            subject=f"【工作紙】{school_name} ({grade}) - {student_name} 的校本填充練習",
            html_content=f"""
                <p>親愛的家長您好：</p>
                <p>附件為 <strong>{student_name}</strong> 同學在 <strong>{school_name} ({grade})</strong> 的校本填充工作紙。</p>
                <p>請下載並列印供同學練習。祝 學習愉快！</p>
                <br><p>-- 自動發送系統 --</p>
            """
        )

        if cc_email:
            cc_clean = str(cc_email).strip().lower()
            if cc_clean not in ["n/a", "nan", "", "none"] and "@" in cc_clean and cc_clean != recipient.lower():
                message.add_cc(cc_clean)

        encoded_pdf = base64.b64encode(pdf_bytes).decode()
        attachment = Attachment(
            FileContent(encoded_pdf),
            FileName(f"{safe_name}_Worksheet.pdf"),
            FileType("application/pdf"),
            Disposition("attachment")
        )
        message.add_attachment(attachment)

        sg = SendGridAPIClient(sg_config["api_key"])
        response = sg.send(message)

        if 200 <= response.status_code < 300:
            return True, "發送成功"
        else:
            return False, f"SendGrid Error: {response.status_code}"

    except HTTPError as e:
        try:
            return False, e.body.decode("utf-8")
        except Exception:
            return False, str(e)

    except Exception as e:
        return False, str(e)

# ============================================================
# --- PDF Preview Helper ---
# ============================================================

def display_pdf_as_images(pdf_bytes):
    try:
        images = convert_from_bytes(pdf_bytes, dpi=150)
        for i, image in enumerate(images):
            st.image(image, caption=f"Page {i+1}", use_container_width=True)
    except Exception as e:
        st.error(f"無法顯示 PDF 預覽: {e}")
        st.info("你仍然可以使用下載按鈕下載 PDF。")

# ============================================================
# --- Load Data ---
# ============================================================

with st.spinner("正在載入資料，請稍候..."):
    student_df = load_students()
    standby_df = load_standby()
    standby_groups = parse_standby_table(standby_df)

# ============================================================
# --- Sidebar Controls ---
# ============================================================

with st.sidebar:
    st.header("⚙️ 控制面板")

    # === 控制區塊 ===
    with st.container(border=True):
        col_r, col_s = st.columns(2)

        with col_r:
            if st.button("🔄 更新資料", use_container_width=True, help="點擊重新載入 Google Sheets 資料"):
                with st.spinner("正在同步最新資料..."):
                    load_students.clear()
                    load_standby.clear()
                    st.session_state.final_pool = {}
                    st.session_state.confirmed_batches = set()
                    st.session_state.shuffled_cache = {}
                    st.rerun()

        with col_s:
            if st.button("🔀 打亂題目", use_container_width=True, help="重新隨機排序題目順序"):
                st.session_state.shuffled_cache = {}
                st.rerun()

    st.divider()

    # === 篩選區塊 ===
    with st.container(border=True):
        # --- 替換開始 (原本的 479-491 行) ---
        st.subheader("🔍 篩選條件")
        
        # 1. 先選學校
        all_schools = sorted({k.split("||")[0] for k in standby_groups}) if standby_groups else ["無資料"]
        selected_school = st.selectbox("🏫 選擇學校", all_schools)
        
        # 2. 根據學校過濾年級
        available_levels = sorted({
            k.split("||")[1] for k in standby_groups 
            if k.startswith(f"{selected_school}||")
        })
        selected_level = st.selectbox(
            "🎓 選擇年級", 
            available_levels if available_levels else ["P1"],
            label_visibility="visible"
        )

        # 3. 組合目前的 Batch Key
        current_batch_key = f"{selected_school}||{selected_level}"

        # 4. 狀態重置邏輯
        if st.session_state.last_selected_level != current_batch_key:
            st.session_state.last_selected_level = current_batch_key
            st.session_state.selected_student_name_b = None
        # --- 替換結束 ---
            st.session_state.selected_student_name_b = None

    st.divider()

    # === 狀態儀表板 ===
    with st.container(border=True):
        st.subheader("📊 資料概覽")

        level_batches = [k for k in standby_groups if k.endswith(f"||{selected_level}")]
        total_words = sum(len(v) for k, v in standby_groups.items() if k.endswith(f"||{selected_level}"))

        # 計算已使用
        if standby_df is not None and not standby_df.empty:
            used_count = standby_df[standby_df["Status"].str.strip() == "已使用"].shape[0]
        else:
            used_count = 0

        available_count = total_words

        confirmed_count = len([k for k in st.session_state.confirmed_batches if k.endswith(f"||{selected_level}")])
        pool_count = sum(
            len(v) for k, v in st.session_state.final_pool.items()
            if k.endswith(f"||{selected_level}") and isinstance(v, list)
        )

        col_stat1, col_stat2 = st.columns(2)
        with col_stat1:
            st.metric("批次數", len(level_batches))
            st.metric("可用詞語", available_count, delta="📝 可用" if available_count > 0 else None)
        with col_stat2:
            st.metric("總詞語", total_words)
            st.metric("已使用", used_count, delta="✅ 已使用" if used_count > 0 else None)

        st.metric("已鎖定題庫", pool_count)

        if not student_df.empty and "狀態" in student_df.columns:
            active_count = (student_df["狀態"] == "Y").sum()
            st.metric("啟用學生", int(active_count))

    st.divider()

    # === 說明區塊 ===
    with st.expander("📖 使用說明", expanded=False):
        st.markdown("""
        **操作流程：**

        1. **鎖定題庫**：確認 Standby 句子後鎖定題目
        2. **預覽下載**：生成並下載工作紙 PDF
        3. **寄送郵件**：將工作紙寄送給學生家長

        **小提示：**
        - 句子格式使用 ＿＿＿＿ 標記填空位置
        - 寄送前請確認學生資料正確
        """)

    # === 系統狀態 ===
    with st.container(border=True):
        st.caption("🔗 系統狀態")
        if not student_df.empty:
            st.success("✅ Google Sheets 已連接")
        else:
            st.warning("⚠️ 請檢查資料連接")

# ============================================================
# --- Shuffle Helper ---
# ============================================================

def get_shuffled_questions(questions, cache_key):
    if cache_key in st.session_state.shuffled_cache:
        return st.session_state.shuffled_cache[cache_key]
    questions_list = list(questions)
    random.seed(int(time.time() * 1000))
    random.shuffle(questions_list)
    st.session_state.shuffled_cache[cache_key] = questions_list
    return questions_list

# ============================================================
# --- PDF Layout Constants ---
# ============================================================

PDF_LEFT_NUM = 60
PDF_TEXT_START = PDF_LEFT_NUM + 30
PDF_RIGHT_MARGIN = 40
PDF_LINE_HEIGHT = 26
PDF_FONT_SIZE = 18

# ============================================================
# --- 頂部標籤頁導航 ---
# ============================================================

st.divider()

tab_lock, tab_preview, tab_email = st.tabs([
    "📥 題庫鎖定（Standby）",
    "📄 預覽下載",
    "✉️ 寄送郵件"
])

# ============================================================
# --- 標籤頁 1: 題庫鎖定 ---
# ============================================================

with tab_lock:
    st.subheader("📥 題庫鎖定（Standby）")

    level_groups = {k: v for k, v in standby_groups.items() if k.endswith(f"||{selected_level}")}

    if not level_groups:
        with st.container(border=True):
            st.success(f"✅ {selected_level} 目前沒有任何可用題目。")
            st.info("請確認 Google Sheets 中的 standby 工作表是否有 Status 為 Ready 的資料，或嘗試點擊側邊欄的「更新資料」按鈕。")
        st.stop()

    for batch_key, word_dict in level_groups.items():
        with st.container(border=True):
            school, level = batch_key.split("||")
            st.markdown(f"### 🏫 {school}（{level}）")

            ready_words, pending_words, is_ready = compute_batch_readiness(batch_key, word_dict)

            with st.expander(f"📝 可用詞語（共 {len(word_dict)} 個）", expanded=True):
                for word, data in word_dict.items():
                    st.markdown(f"- **{word}**: {data.get('content', '')}")

            if is_ready and batch_key not in st.session_state.confirmed_batches:
                with st.container(border=True):
                    st.markdown("### 🔒 確認並鎖定題庫")

                    row_indices = [
                        data["row_index"]
                        for data in word_dict.values()
                        if "row_index" in data
                    ]

                    st.info(f"即將鎖定並標記 {len(row_indices)} 個句子為「已使用」。")

                    confirm_checkbox = st.checkbox(
                        "我確認要鎖定題庫並將這些句子標記為已使用",
                        key=f"confirm_check_{batch_key}"
                    )

                    if confirm_checkbox:
                        if st.button(f"✅ 確認並鎖定題庫：{school}", key=f"confirm_{batch_key}", type="primary"):
                            with st.spinner("正在鎖定題庫並更新 Status..."):
                                final_qs = build_final_pool_for_batch(batch_key, word_dict)
                                st.session_state.final_pool[batch_key] = final_qs
                                st.session_state.confirmed_batches.add(batch_key)

                                if row_indices:
                                    update_ok, update_msg = update_status_to_used(row_indices)
                                    if update_ok:
                                        st.success(f"✅ 已成功鎖定題庫並更新 {len(row_indices)} 個句子的 Status")
                                    else:
                                        st.error(f"❌ 更新失敗：{update_msg}")
                                        st.info("💡 請確保 Google Service Account 有試算表的編輯權限")

                            st.rerun()
                    else:
                        st.caption("請勾選上方確認方塊以啟用鎖定按鈕")

            elif batch_key in st.session_state.confirmed_batches:
                st.success("✅ 此批次已完成並已標記為已使用。")

# ============================================================
# --- 標籤頁 2: 預覽下載 ---
# ============================================================

with tab_preview:
    st.subheader("📄 預覽下載")

    level_batches = {k: v for k, v in st.session_state.final_pool.items() if k.endswith(f"||{selected_level}")}

    if not level_batches:
        with st.container(border=True):
            st.warning("⚠️ 尚未有任何批次完成鎖定題庫。")
            st.info("請先到「題庫鎖定」標籤頁完成鎖定後，再回到此處下載工作紙。")
        st.stop()

    for batch_key, questions in level_batches.items():
        with st.container(border=True):
            school, level = batch_key.split("||")
            st.markdown(f"### 🏫 {school}（{level}）")
            st.caption(f"共 {len(questions)} 題")

            # --- 新增：隨機排序邏輯 ---
            # 使用 batch_key 作為快取鍵，確保同一個批次在本次 Session 中順序固定，但點擊側邊欄「打亂題目」會更新
            shuffled_qs = get_shuffled_questions(questions, f"preview_{batch_key}")

            with st.spinner("正在生成 PDF..."):
                # 使用隨機排序後的 shuffled_qs 生成 PDF
                pdf_bytes = create_pdf(school, level, shuffled_qs)
                answer_pdf_bytes = create_answer_pdf(school, level, shuffled_qs)

            col1, col2 = st.columns(2)

            with col1:
                st.download_button(
                    label="⬇️ 下載學生版 PDF",
                    data=pdf_bytes,
                    file_name=f"{school}_{level}_worksheet.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    help="下載學生版本的工作紙 PDF"
                )

            with col2:
                st.download_button(
                    label="⬇️ 下載教師版 PDF（答案）",
                    data=answer_pdf_bytes,
                    file_name=f"{school}_{level}_answers.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    help="下載包含答案的教師版 PDF"
                )

            with st.expander("📘 預覽學生版 PDF", expanded=False):
                display_pdf_as_images(pdf_bytes)

# ============================================================
# --- 標籤頁 3: 寄送郵件 ---
# ============================================================

with tab_email:
    st.subheader("✉️ 寄送郵件")

    if student_df.empty:
        st.error("❌ 學生資料表為空，無法寄送。")
        st.stop()

    # --- 優化點 1：聯動篩選 ---
    # 根據側邊欄選中的「學校」和「年級」精確過濾學生名單
    df_filtered = student_df[
        (student_df["學校"].astype(str) == selected_school) & 
        (student_df["年級"].astype(str) == selected_level)
    ]

    if df_filtered.empty:
        with st.container(border=True):
            st.warning(f"⚠️ 在 {selected_school} 的 {selected_level} 年級中找不到學生資料。")
            st.info("請確認「學生資料」工作表中的學校名稱與年級是否完全匹配。")
        st.stop()

    # --- 優化點 2：顯示過濾後的名單 ---
    with st.container(border=True):
        st.markdown(f"### 👤 選擇學生 ({selected_school} - {selected_level})")
        
        # 排序學生姓名，讓找人更直覺
        student_names = sorted(df_filtered["學生姓名"].tolist())
        
        selected_student = st.selectbox(
            "請輸入或選擇學生姓名",
            [""] + student_names,
            help="提示：點擊後直接輸入姓名可快速搜尋",
            key="student_selector_main"
        )

    if not selected_student:
        st.info("👆 請從上方選擇一位學生以開始寄送流程")
        st.stop()

    # 獲取選中學生的詳細資料
    row = df_filtered[df_filtered["學生姓名"] == selected_student].iloc[0]
    # ... (後續的 PDF 生成與寄送邏輯保持不變)
    school = row["學校"]
    grade = row["年級"]
    parent_email = row.get("家長 Email", "")
    cc_email = row.get("老師 Email", "")

    batch_key = f"{school}||{grade}"

    if batch_key not in st.session_state.final_pool:
        with st.container(border=True):
            st.error("⚠️ 此學生所屬批次尚未完成鎖定題庫。")
            st.info("請先到「題庫鎖定」標籤頁完成鎖定。")
        st.stop()

    questions = st.session_state.final_pool[batch_key]

    with st.container(border=True):
        st.markdown("### 📄 工作紙預覽")

        with st.spinner("正在生成 PDF..."):
            # 這裡加上 .getvalue() 把文件對象轉成純數據
            pdf_obj = create_pdf(school, grade, questions, student_name=selected_student)
            pdf_bytes = pdf_obj.getvalue()

        st.download_button(
            label="⬇️ 下載學生版 PDF",
            data=pdf_bytes,
            file_name=f"{selected_student}_worksheet.pdf",
            mime="application/pdf",
            use_container_width=True
        )

    st.divider()

    with st.container(border=True):
        st.markdown("### ✉️ 寄送工作紙")

        with st.expander("📋 寄送資訊摘要", expanded=True):
            st.markdown(f"""
            - **學生姓名**：{selected_student}
            - **學校**：{school}
            - **年級**：{grade}
            - **家長電郵**：{parent_email if parent_email else '（未提供）'}
            - **老師電郵**：{cc_email if cc_email else '（未提供）'}
            """)

        st.markdown("#### ⚠️ 確認寄送")

        if not parent_email or parent_email.lower() in ["n/a", "nan", "", "none"]:
            st.error("❌ 該學生的家長電郵地址為空，無法寄送。")
            st.stop()

        confirm_email = st.checkbox(
            f"我確認要將工作紙寄送至以下電郵：{parent_email}",
            key="email_confirm_checkbox"
        )

        if not confirm_email:
            st.caption("請勾選上方確認方塊以啟用寄送按鈕")
            st.stop()

        if st.button("📨 寄出工作紙", type="primary", use_container_width=True):
            with st.spinner("正在發送郵件，請稍候..."):
                ok, msg = send_email_with_pdf(
                    parent_email,
                    selected_student,
                    school,
                    grade,
                    pdf_bytes,
                    cc_email=cc_email
                )

            if ok:
                st.success("🎉 已成功寄出工作紙！")
                st.balloons()
                st.toast(f"工作紙已成功寄送給 {selected_student} 的家長！", icon="✅")
            else:
                st.error(f"❌ 寄送失敗：{msg}")
                st.info("請檢查網路連線或稍後再試。")

# ============================================================
# --- End of App ---
# ============================================================

st.write("")
st.write("© 2026 校本填充工作紙生成器 — 自動化教學工具")
